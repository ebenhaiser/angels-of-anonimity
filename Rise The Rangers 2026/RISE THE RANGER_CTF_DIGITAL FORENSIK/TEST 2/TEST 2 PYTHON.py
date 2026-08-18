from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import hashlib, os

base = Path("/mnt/data")
files = [
    "boarding_pass.pdf",
    "shipping_note.pdf",
    "4983789463483.db",
    "4983789463483 (1).db",
    "foremost.7z",
    "jpseek.7z",
    "steghide.7z",
]

def sha256(path):
    h=hashlib.sha256()
    with open(path,"rb") as f:
        for chunk in iter(lambda:f.read(1024*1024), b""):
            h.update(chunk)
    return h.hexdigest()

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("LAPORAN DOKUMENTASI FORENSIK DIGITAL")
r.bold = True
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Investigasi Artefak Chapter 1 dan Chapter 2")
r.bold = True
r.font.size = Pt(13)

doc.add_paragraph("")

def heading(text, level=1):
    doc.add_heading(text, level=level)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

heading("1. Tujuan Investigasi")
doc.add_paragraph(
    "Dokumentasi ini mencatat proses pemeriksaan forensik terhadap artefak digital "
    "yang digunakan pada challenge Chapter 1 dan Chapter 2. Fokus pemeriksaan adalah "
    "identifikasi artefak, preservasi bukti, pemeriksaan metadata dan isi dokumen, "
    "serta analisis indikasi steganografi. Nilai yang belum dapat dibuktikan dari "
    "artefak yang tersedia dicatat sebagai belum terkonfirmasi dan tidak diisi dengan tebakan."
)

heading("2. Prinsip dan Metodologi")
for x in [
    "Preservasi: artefak dianalisis tanpa mengubah file sumber.",
    "Identifikasi: nama file, ukuran, tipe file, dan hash SHA-256 dicatat.",
    "Pemeriksaan statis: isi PDF, metadata PDF, dan teks yang dapat diekstrak diperiksa.",
    "Pemeriksaan steganografi: artefak JPEG/Steghide/JPSeek dicatat sebagai jalur analisis terpisah.",
    "Korelasi: temuan antar-node dibandingkan untuk membentuk rangkaian informasi.",
    "Pelaporan: setiap kesimpulan dibedakan antara temuan langsung, hasil ekstraksi, dan hal yang belum terverifikasi."
]:
    bullet(x)

heading("3. Inventaris Barang Bukti")
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Artefak"
hdr[1].text = "Ukuran"
hdr[2].text = "SHA-256"
hdr[3].text = "Peran dalam analisis"
roles = {
    "boarding_pass.pdf":"Boarding pass / metadata PDF",
    "shipping_note.pdf":"Hasil ekstraksi Steghide",
    "4983789463483.db":"Database",
    "4983789463483 (1).db":"Salinan database identik",
    "foremost.7z":"Hasil carving Foremost",
    "jpseek.7z":"Artefak analisis JPSeek",
    "steghide.7z":"Artefak hasil Steghide",
}
for fn in files:
    path=base/fn
    cells=table.add_row().cells
    cells[0].text=fn
    cells[1].text=f"{path.stat().st_size:,} bytes"
    cells[2].text=sha256(path)
    cells[3].text=roles[fn]

heading("4. Chapter 1 – Temuan")
t = doc.add_table(rows=1, cols=3)
t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,v in enumerate(["Node","Jawaban","Status"]): t.rows[0].cells[i].text=v
ch1 = [
    ("Node 1","DANIEL KOVACS","Teridentifikasi"),
    ("Node 2","KRANSTADT","Teridentifikasi"),
    ("Node 3","PT KIRANA TRANS LOGISTIK","Teridentifikasi dari rangkaian artefak"),
    ("Node 4","CANGGU","Teridentifikasi dari rangkaian artefak"),
]
for row in ch1:
    c=t.add_row().cells
    for i,v in enumerate(row): c[i].text=v

heading("5. Chapter 2 – Temuan")
t = doc.add_table(rows=1, cols=3)
t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,v in enumerate(["Node","Jawaban","Status"]): t.rows[0].cells[i].text=v
ch2 = [
    ("Node 01","0.15 BTC","Jawaban yang diberikan pada challenge"),
    ("Node 02","PAK ANTON","Jawaban yang diberikan pada challenge; digunakan sebagai password Steghide"),
    ("Node 03","BELUM TERKONFIRMASI","Pertanyaan asli Node 03 tidak tersedia dalam bentuk teks pada artefak yang dapat dibaca saat penyusunan laporan"),
]
for row in ch2:
    c=t.add_row().cells
    for i,v in enumerate(row): c[i].text=v

heading("6. Pemeriksaan Boarding Pass")
doc.add_paragraph(
    "File boarding_pass.pdf diperiksa pada level isi dan metadata. Teks halaman menunjukkan "
    "nama penumpang KOVACS / DANIEL MR, rute SINGAPORE (SIN) ke DENPASAR - BALI (DPS), "
    "penerbangan GA-9147, tanggal 01 JUN 2026, gate D7, seat 14C, boarding 13:20, dan sequence 0142."
)
doc.add_paragraph("Metadata PDF yang teridentifikasi:")
for x in [
    "Author: PT KIRANA TRANS LOGISTIK",
    "Creator: Garuda Island Air E-Ticketing",
    "Producer: PT KIRANA TRANS LOGISTIK - Document Processing System v2.1",
    "CreationDate: 01 Agustus 2026 15:22:30 UTC",
]:
    bullet(x)
doc.add_paragraph(
    "Catatan forensik: field Author, Creator, dan Producer adalah field metadata yang berbeda. "
    "Karena itu, kesimpulan tentang 'pembuat tiket' harus mengikuti redaksi pertanyaan challenge "
    "dan field yang memang diminta; metadata saja tidak cukup untuk menyamakan ketiganya."
)

heading("7. Pemeriksaan Steghide dan Korelasi Password")
doc.add_paragraph(
    "Password PAKANTON digunakan pada proses Steghide. Artefak hasil proses tersebut kemudian "
    "menghasilkan file shipping_note.pdf. Hal ini merupakan bukti penting bahwa PAKANTON bukan "
    "sekadar kandidat password: password tersebut berhasil digunakan pada tahap ekstraksi yang dilaporkan."
)
doc.add_paragraph("Rantai artefak:")
for x in [
    "Chapter 2 Node 02 → PAK ANTON",
    "Password normalisasi yang dicoba → PAKANTON",
    "Proses Steghide → menghasilkan shipping_note.pdf",
]:
    bullet(x)

heading("8. Analisis shipping_note.pdf")
doc.add_paragraph(
    "shipping_note.pdf berisi nota pengiriman dari EXPRESS CARGO NUSANTARA. Informasi yang dapat "
    "diverifikasi dari teks dokumen:"
)
for x in [
    "No. Invoice: EXP-2026-06-0847",
    "Tanggal pemesanan: 13 Juni 2026",
    "Pengirim: Toko Elektronik Sumber Jaya - Denpasar",
    "Penerima: A. Wisnu (PT Cahaya Nusantara Energi)",
    "Jenis barang: Komponen Elektronik (1 paket, 2.4 kg)",
    "Kode referensi: KTL-88291X",
    "Metode: Ekspedisi Reguler - Jalur Darat/Penyeberangan",
    "Estimasi tiba: 18 Juni 2026",
    "Catatan: barang diteruskan melalui jalur penyeberangan reguler dan mengikuti jadwal kapal feri harian",
]:
    bullet(x)
doc.add_paragraph("Metadata PDF:")
for x in [
    "Author: anonymous",
    "Creator: anonymous",
    "Producer: ReportLab PDF Library - (opensource)",
    "Keywords: GILIMANUK, internal-routing, non-standard-channel",
    "Subject: Internal routing note: shipment KTL-88291X cleared via Pelabuhan Gilimanuk crossing point, non-standard freight channel",
    "CreationDate / ModDate: 02 Agustus 2026 01:17:54 UTC",
]:
    bullet(x)
doc.add_paragraph(
    "Temuan korelasi: GILIMANUK muncul langsung pada Keywords dan Subject metadata, sedangkan "
    "Pelabuhan Gilimanuk disebut pada Subject. Namun teks isi nota hanya menyebut jalur penyeberangan "
    "reguler dan tidak menyebut nama pelabuhan secara eksplisit. Karena itu, GILIMANUK merupakan temuan "
    "metadata/subject, bukan hasil pembacaan field isi utama."
)

heading("9. Pemeriksaan Database dan Artefak Carving")
doc.add_paragraph(
    "Dua file database 4983789463483.db dan 4983789463483 (1).db memiliki ukuran yang sama (16,384 bytes) "
    "dan SHA-256 yang sama. Ini menunjukkan keduanya identik secara byte pada saat pemeriksaan."
)
doc.add_paragraph(
    "Artefak foremost.7z dan jpseek.7z dicatat sebagai hasil/alat analisis. Kehadiran archive tersebut "
    "tidak dengan sendirinya membuktikan isi payload tertentu; hasil ekstraksi perlu dikaitkan dengan "
    "file sumber dan hash untuk menjadi temuan yang dapat dipertanggungjawabkan."
)

heading("10. Status Node 03 Chapter 2")
doc.add_paragraph(
    "Node 03 belum dapat ditetapkan secara forensik dari bahan yang tersedia dalam sesi ini karena "
    "teks pertanyaan Node 03 tidak terbaca/tersedia sebagai sumber yang dapat dikutip. Beberapa jawaban "
    "sempat diasumsikan sebelumnya, tetapi asumsi tersebut tidak boleh dimasukkan sebagai finding final."
)
doc.add_paragraph(
    "Evidence yang tersedia untuk melanjutkan Node 03 adalah hasil Steghide berupa shipping_note.pdf "
    "beserta metadata dan isi dokumennya. Pertanyaan asli Node 03 diperlukan untuk menentukan field "
    "mana yang harus dijadikan jawaban."
)

heading("11. Kesimpulan Forensik")
for x in [
    "Rangkaian artefak Chapter 1 yang telah dicatat: DANIEL KOVACS, KRANSTADT, PT KIRANA TRANS LOGISTIK, CANGGU.",
    "Rangkaian Chapter 2 yang terkonfirmasi dari informasi challenge: 0.15 BTC dan PAK ANTON.",
    "PAKANTON berhasil digunakan pada tahap Steghide dan menghasilkan shipping_note.pdf.",
    "shipping_note.pdf mengandung kode referensi KTL-88291X serta metadata yang menunjuk pada GILIMANUK dan routing internal/non-standard channel.",
    "Node 03 Chapter 2 tetap berstatus pending sampai pertanyaan asli Node 03 tersedia dan dapat dikorelasikan dengan evidence."
]:
    bullet(x)

heading("12. Catatan Integritas Bukti")
doc.add_paragraph(
    "Hash SHA-256 pada bagian inventaris digunakan sebagai fingerprint integritas file yang diperiksa. "
    "Untuk kebutuhan laporan formal/penyerahan bukti, proses akuisisi sebaiknya dilengkapi timestamp, "
    "identitas pemeriksa, sumber media, metode akuisisi, tool dan versinya, serta hash sebelum dan sesudah "
    "setiap tahap ekstraksi."
)

out = base/"Dokumentasi_Forensik_Chapter_1_2.docx"
doc.save(out)
print(out)
