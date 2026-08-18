from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import hashlib, os, datetime

out = "/mnt/data/Laporan_Forensik_Digital_Chapter_5.docx"

# Evidence files available in the session
evidence = [
    ("/mnt/data/777766767676.jpeg", "Artifact foto JPEG – Node 3"),
    ("/mnt/data/4364836483748374.csv", "Daftar personel – Node 1"),
    ("/mnt/data/975486584754.txt", "Internal operation registry – Node 2"),
    ("/mnt/data/foremost.7z", "Hasil carving/foremost – Node 3"),
    ("/mnt/data/jpseek.7z", "Hasil analisis JPEG – Node 3"),
]
def sha256(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024*1024), b""):
            h.update(chunk)
    return h.hexdigest()

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(20)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(15)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN FORENSIK DIGITAL")
r.bold = True
r.font.size = Pt(20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CHAPTER 5 – INVESTIGASI MULTI-ARTIFACT")
r.bold = True
r.font.size = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Dokumentasi analisis Node 1–Node 5").italic = True

doc.add_paragraph("Dokumen ini mendokumentasikan proses pemeriksaan artifact, korelasi lintas-node, hasil analisis, dan chain of custody untuk Chapter 5. Isi laporan membedakan fakta yang diperoleh dari artifact dengan kesimpulan/hasil flag yang telah dikonfirmasi pada pengerjaan challenge.")

doc.add_heading("1. Identitas Pemeriksaan", level=1)
tbl = doc.add_table(rows=0, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for a,b in [
    ("Kasus", "Chapter 5 – Digital Forensics / CTF"),
    ("Ruang lingkup", "Node 1 sampai Node 5"),
    ("Metodologi", "Preservasi artifact, pemeriksaan metadata, decoding, korelasi lintas-artifact, validasi flag"),
    ("Tanggal dokumentasi", "15 Agustus 2026"),
    ("Status", "Selesai – seluruh node terkonfirmasi"),
]:
    cells = tbl.add_row().cells
    cells[0].text = a
    cells[1].text = b

doc.add_heading("2. Tujuan Pemeriksaan", level=1)
for t in [
    "Mengidentifikasi target personel pada Node 1 dari daftar personel.",
    "Menghubungkan kode operasi internal dengan alias kelompok pada Node 2.",
    "Menentukan lokasi pada Node 3 melalui pemeriksaan artifact foto dan korelasi lokasi.",
    "Menentukan waktu kejadian pada Node 4 melalui decoding petunjuk dan korelasi jadwal acara.",
    "Menghasilkan identifikasi akhir berupa kombinasi kelompok dan insider pada Node 5."
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("3. Daftar Evidence dan Hash", level=1)
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i,t in enumerate(["Evidence", "Fungsi", "SHA-256", "Status"]):
    hdr[i].text = t
for path, desc in evidence:
    row = tbl.add_row().cells
    row[0].text = os.path.basename(path)
    row[1].text = desc
    row[2].text = sha256(path) if os.path.exists(path) else "Tidak tersedia pada runtime"
    row[3].text = "Diperiksa" if os.path.exists(path) else "Tidak tersedia"

doc.add_paragraph("Catatan chain of custody: hash di atas adalah hash file yang tersedia pada runtime saat penyusunan laporan. Waktu akuisisi asli, operator akuisisi, media sumber, dan serial number perangkat tidak tersedia dalam data yang diberikan sehingga tidak diisi secara asumtif.")

doc.add_heading("4. Metodologi Forensik", level=1)
for t in [
    "Preservasi: artifact diperlakukan sebagai evidence dan tidak dilakukan perubahan terhadap file sumber.",
    "Identifikasi: nama file, tipe file, ukuran, dan struktur dasar diperiksa.",
    "Pemeriksaan metadata: khusus artifact JPEG, EXIF dan GPS dianalisis.",
    "Analisis isi: CSV diperiksa untuk korelasi personel; registry diperiksa untuk pemetaan kode operasi; ciphertext dianalisis menggunakan Caesar shift.",
    "Korelasi: hasil setiap node dibandingkan dengan artifact dan hasil node sebelumnya.",
    "Validasi: jawaban akhir dicatat berdasarkan hasil yang dikonfirmasi pada challenge."
]:
    doc.add_paragraph(t, style="List Number")

doc.add_heading("5. Analisis Node 1 – Identifikasi Insider", level=1)
doc.add_paragraph("Artifact: 4364836483748374.csv")
doc.add_paragraph("CSV berisi 41 record personel dengan kolom employee_id, nama_lengkap, departemen, posisi, tanggal_bergabung, dan status. Seluruh record yang diperiksa berstatus AKTIF.")
doc.add_paragraph("Hasil korelasi pada challenge menunjuk kepada personel berikut:")
tbl = doc.add_table(rows=1, cols=3)
for i,t in enumerate(["Field", "Hasil", "Keterangan"]):
    tbl.rows[0].cells[i].text = t
for a,b,c in [
    ("Nama lengkap", "ARYA WISNU NUGRAHA", "Jawaban Node 1 terkonfirmasi"),
    ("Employee ID", "KRY-2291", "Record pada CSV"),
    ("Departemen", "Teknisi SCADA", "Record pada CSV"),
    ("Posisi", "Teknisi Senior", "Record pada CSV"),
    ("Tanggal bergabung", "14-03-2021", "Record pada CSV"),
    ("Status", "AKTIF", "Record pada CSV"),
]:
    cells=tbl.add_row().cells
    cells[0].text=a; cells[1].text=b; cells[2].text=c

doc.add_heading("6. Analisis Node 2 – Identifikasi Kelompok", level=1)
doc.add_paragraph("Artifact: 975486584754.txt (internal_ops_registry.dump)")
doc.add_paragraph("Registry menggunakan format internal_op_code => group_alias. Salah satu entri yang relevan adalah NW-7739 => NIGHT WOLF. Registry juga menjelaskan bahwa kode operasi bersifat unik per kelompok pada periode aktif dan beberapa kelompok dapat memiliki lebih dari satu kode.")
doc.add_paragraph("Hasil Node 2 yang telah dikonfirmasi: NIGHT WOLF.")
doc.add_paragraph("Korelasi: kode operasi internal yang relevan dipetakan ke alias kelompok NIGHT WOLF berdasarkan registry.")

doc.add_heading("7. Analisis Node 3 – Identifikasi Lokasi", level=1)
doc.add_paragraph("Artifact utama: 777766767676.jpeg, dengan hasil analisis tambahan dari foremost.7z dan jpseek.7z.")
doc.add_paragraph("Pemeriksaan EXIF pada JPEG menunjukkan informasi GPS dan waktu berikut:")
tbl = doc.add_table(rows=1, cols=2)
tbl.rows[0].cells[0].text="Parameter"
tbl.rows[0].cells[1].text="Nilai"
for a,b in [
    ("Latitude", "8°19′06.6″ S"),
    ("Longitude", "114°53′32.64″ E"),
    ("Koordinat desimal", "-8.3185, 114.8924"),
    ("Altitude", "650 m"),
    ("Timestamp EXIF", "2026:06:19 18:05:42"),
]:
    cells=tbl.add_row().cells
    cells[0].text=a; cells[1].text=b

doc.add_paragraph("Tahap awal sempat menghasilkan beberapa kandidat lokasi administratif yang tidak diterima oleh challenge. Setelah korelasi ulang artifact, hasil Node 3 yang dikonfirmasi adalah SEPANG KELOD. Dengan demikian, nama lokasi tersebut dicatat sebagai finding final, bukan sekadar label locality terdekat dari koordinat.")
doc.add_paragraph("Hasil Node 3: SEPANG KELOD.")

doc.add_heading("8. Analisis Node 4 – Penentuan Waktu", level=1)
doc.add_paragraph("Ciphertext yang diperiksa: “WLJD PHQLV WHEORXP SXQFDN”.")
doc.add_paragraph("Pemeriksaan Caesar cipher dengan pergeseran -3 menghasilkan teks yang secara semantik menunjuk pada petunjuk “TIGA MENIT SEBELUM PUNCAK” (ciphertext memiliki ketidaktepatan karakter sehingga hasil literal tidak seluruhnya sempurna).")
doc.add_paragraph("Korelasi dengan evidence acara menunjukkan waktu puncak yang relevan adalah 21:00 WITA. Tiga menit sebelum waktu tersebut adalah 20:57 WITA, pada tanggal 20 Juni 2026.")
doc.add_paragraph("Hasil Node 4 yang telah dikonfirmasi: 20-06-2026 20:57.")

doc.add_heading("9. Analisis Node 5 – Final Identification", level=1)
doc.add_paragraph("Node 5 meminta kombinasi identitas kelompok dan insider yang telah dikonfirmasi pada node sebelumnya.")
doc.add_paragraph("Korelasi:")
for t in [
    "Node 1 → ARYA WISNU NUGRAHA",
    "Node 2 → NIGHT WOLF",
    "Node 3 → SEPANG KELOD",
    "Node 4 → 20-06-2026 20:57"
]:
    doc.add_paragraph(t, style="List Bullet")
doc.add_paragraph("Final identification yang dikonfirmasi: NIGHT WOLF - ARYA WISNU NUGRAHA.")

doc.add_heading("10. Chain of Custody", level=1)
tbl = doc.add_table(rows=1, cols=6)
for i,t in enumerate(["No.", "Evidence", "Diterima", "Proses", "Hash SHA-256", "Catatan"]):
    tbl.rows[0].cells[i].text=t
for idx,(path,desc) in enumerate(evidence,1):
    cells=tbl.add_row().cells
    cells[0].text=str(idx)
    cells[1].text=os.path.basename(path)
    cells[2].text="Artifact diunggah ke sesi pemeriksaan"
    cells[3].text="Analisis forensik / korelasi challenge"
    cells[4].text=sha256(path) if os.path.exists(path) else "N/A"
    cells[5].text="Integrity hash dicatat saat dokumentasi"

doc.add_paragraph("Batasan chain of custody: dokumentasi ini merekam alur evidence dalam lingkungan challenge. Tidak tersedia informasi independen mengenai media akuisisi, identitas petugas akuisisi, waktu serah-terima fisik, atau write-blocker. Karena itu elemen tersebut tidak direkonstruksi secara spekulatif.")

doc.add_heading("11. Timeline Temuan", level=1)
tbl=doc.add_table(rows=1, cols=3)
for i,t in enumerate(["Tahap", "Temuan", "Status"]):
    tbl.rows[0].cells[i].text=t
for a,b,c in [
    ("Node 1", "ARYA WISNU NUGRAHA", "TERKONFIRMASI"),
    ("Node 2", "NIGHT WOLF", "TERKONFIRMASI"),
    ("Node 3", "SEPANG KELOD", "TERKONFIRMASI"),
    ("Node 4", "20-06-2026 20:57", "TERKONFIRMASI"),
    ("Node 5", "NIGHT WOLF - ARYA WISNU NUGRAHA", "TERKONFIRMASI"),
]:
    cells=tbl.add_row().cells
    cells[0].text=a; cells[1].text=b; cells[2].text=c

doc.add_heading("12. Kesimpulan", level=1)
doc.add_paragraph("Pemeriksaan Chapter 5 berhasil menghubungkan beberapa artifact berbeda melalui analisis data tabular, registry kode operasi, metadata EXIF/GPS, decoding ciphertext, serta korelasi lintas-node. Seluruh jawaban Node 1 sampai Node 5 telah dikonfirmasi dalam pengerjaan challenge.")
doc.add_paragraph("Final identification: NIGHT WOLF - ARYA WISNU NUGRAHA.")
doc.add_paragraph("Catatan penting: kesimpulan laporan ini mengikuti hasil artifact dan konfirmasi challenge. Informasi yang tidak tersedia pada evidence tidak diisi dengan asumsi.")

doc.add_heading("13. Ringkasan Final Flag", level=1)
p=doc.add_paragraph()
r=p.add_run("NIGHT WOLF - ARYA WISNU NUGRAHA")
r.bold=True
r.font.size=Pt(13)

doc.add_paragraph()
doc.add_paragraph("Dokumen disusun sebagai dokumentasi teknis/forensik untuk keperluan challenge dan pembelajaran digital forensics.")

doc.save(out)
print(out)
print("size:", os.path.getsize(out))
