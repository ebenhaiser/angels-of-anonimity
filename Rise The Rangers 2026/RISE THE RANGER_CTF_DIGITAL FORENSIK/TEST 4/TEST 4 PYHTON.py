from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from datetime import datetime

out = Path("/mnt/data/Dokumentasi_Forensik_Chapter_4.docx")

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(20)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(14)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(11.5)

def shade_cell(cell, fill="D9EAF7"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN DOKUMENTASI FORENSIK DIGITAL")
r.bold = True
r.font.size = Pt(20)
r.font.name = "Arial"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CHAPTER 4 — ANALISIS NODE 01–04")
r.bold = True
r.font.size = Pt(15)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("\nDigital Forensics / CTF Investigation Report").italic = True

doc.add_paragraph()
meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = "Table Grid"
for i, (k,v) in enumerate([
    ("Jenis pemeriksaan", "Digital Forensic Artifact Analysis"),
    ("Ruang lingkup", "Node 01, Node 02, Node 03, Node 04"),
    ("Tujuan", "Ekstraksi artefak, korelasi bukti, validasi jawaban node"),
    ("Status", "Selesai / hasil terverifikasi terhadap challenge"),
    ("Tanggal dokumentasi", datetime.now().strftime("%d-%m-%Y")),
]):
    set_cell_text(meta.cell(i,0), k, True)
    set_cell_text(meta.cell(i,1), v)

doc.add_page_break()

doc.add_heading("1. Ringkasan Eksekutif", level=1)
doc.add_paragraph(
    "Pemeriksaan ini mendokumentasikan proses analisis forensik digital pada empat node "
    "dalam Chapter 4. Pendekatan yang digunakan meliputi pemeriksaan artefak teks, "
    "korelasi transaksi blockchain, decoding Base64, analisis XOR, korelasi DNS/PCAP, "
    "serta pemeriksaan metadata EXIF/GPS pada citra."
)
doc.add_paragraph(
    "Hasil akhir yang diperoleh adalah: Node 01 = 0.45 BTC; Node 02 = "
    "17SkEw2md5avVNyYgj6RiXuQKNwkXaxFyQ; Node 03 = NW-7739|185.220.101.47; "
    "Node 04 = DANAU BERATAN."
)

doc.add_heading("2. Tujuan Pemeriksaan", level=1)
for t in [
    "Mengidentifikasi artefak digital yang relevan pada setiap node.",
    "Menentukan metode decoding/dekripsi yang digunakan oleh artefak.",
    "Menghubungkan artefak antar-node melalui indikator teknis yang konsisten.",
    "Mendokumentasikan evidence, metode, hasil, dan keterbatasan analisis.",
    "Menghasilkan jawaban node berdasarkan evidence, bukan tebakan."
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("3. Metodologi Forensik", level=1)
steps = [
    ("Preservation", "Artefak diperlakukan sebagai evidence dan dianalisis tanpa mengubah isi sumber."),
    ("Identification", "Mengidentifikasi tipe file, field penting, encoding, metadata, dan indikator jaringan."),
    ("Extraction", "Mengekstrak payload Base64, plaintext hasil XOR, field transaksi, DNS, serta EXIF/GPS."),
    ("Correlation", "Mencocokkan indikator antar-artifact, terutama hostname pada beacon dengan DNS answer pada PCAP."),
    ("Validation", "Jawaban akhir hanya ditetapkan setelah indikator teknis saling mendukung."),
    ("Documentation", "Seluruh temuan dirangkum dengan evidence → metode → hasil → kesimpulan.")
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_text(tbl.cell(0,0), "Tahap", True); shade_cell(tbl.cell(0,0))
set_cell_text(tbl.cell(0,1), "Implementasi", True); shade_cell(tbl.cell(0,1))
for a,b in steps:
    cells = tbl.add_row().cells
    set_cell_text(cells[0], a)
    set_cell_text(cells[1], b)

doc.add_heading("4. Analisis Node 01", level=1)
doc.add_paragraph("Temuan utama: nilai transaksi yang menjadi jawaban node adalah 0.45 BTC.")
doc.add_paragraph(
    "Analisis difokuskan pada artefak transaksi/catatan wallet. Nilai nominal diperlakukan "
    "sebagai indikator hasil node, bukan sebagai dasar untuk menebak nilai pada node lain."
)
doc.add_paragraph("Jawaban Node 01:", style=None).runs[0].bold = True
p = doc.add_paragraph()
p.add_run("0.45 BTC").bold = True

doc.add_heading("5. Analisis Node 02", level=1)
doc.add_paragraph(
    "Artefak catatan transaksi dari perangkat cadangan memuat TXID "
    "a1075db55d416d3ca199f55b6084e2115b9345e16c5cf302fc80e9d5fbf5d48d. "
    "Catatan artefak menyatakan transaksi tercatat permanen pada blockchain Bitcoin dan "
    "dapat diverifikasi silang melalui blockchain explorer publik."
)
doc.add_paragraph(
    "Hasil challenge yang telah divalidasi adalah address Bitcoin: "
)
p = doc.add_paragraph()
p.add_run("17SkEw2md5avVNyYgj6RiXuQKNwkXaxFyQ").bold = True

doc.add_heading("6. Analisis Node 03 — Beacon Configuration", level=1)
doc.add_paragraph(
    "Artefak beacon_cfg.dat ditemukan dalam bentuk string Base64. Setelah Base64 decoding, "
    "payload menghasilkan data biner yang tidak langsung terbaca. Analisis known-plaintext "
    "menunjukkan pola XOR berulang dengan key WHISPER_CANGGU."
)
doc.add_paragraph("Plaintext hasil dekripsi:")
code = """OPS-CODE: NW-7739
PRIMARY-UPLINK: relay-sync.nightops-cdn.net
STATUS: ACTIVE
LAST-CHECKIN: 2026-06-18 06:52 WITA"""
p = doc.add_paragraph()
p.style = "No Spacing"
p.add_run(code).font.name = "Courier New"

doc.add_paragraph(
    "Hostname relay-sync.nightops-cdn.net kemudian dikorelasikan dengan artefak PCAP. "
    "DNS response pada PCAP memberikan A record 185.220.101.47. Korelasi ini merupakan "
    "dasar penetapan IP callback, sehingga IP tidak diperoleh melalui brute-force."
)
doc.add_paragraph("Jawaban Node 03:", style=None).runs[0].bold = True
p = doc.add_paragraph()
p.add_run("NW-7739|185.220.101.47").bold = True

doc.add_heading("7. Analisis Node 04 — Image Metadata", level=1)
doc.add_paragraph(
    "Citra dianalisis menggunakan metadata EXIF. Field GPS menunjukkan koordinat sekitar "
    "-8.275200, 115.165600. Koordinat tersebut mengarah ke kawasan Danau Beratan, Bali, "
    "yang menjadi dasar penetapan lokasi node."
)
doc.add_paragraph(
    "Validasi challenge menetapkan nama lokasi yang diminta adalah DANAU BERATAN, "
    "bukan nama kawasan yang lebih luas seperti Bedugul."
)
doc.add_paragraph("Jawaban Node 04:", style=None).runs[0].bold = True
p = doc.add_paragraph()
p.add_run("DANAU BERATAN").bold = True

doc.add_heading("8. Korelasi Antar-Node", level=1)
corr = doc.add_table(rows=1, cols=4)
corr.style = "Table Grid"
corr.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Node", "Evidence utama", "Metode", "Hasil"]
for i,h in enumerate(headers):
    set_cell_text(corr.cell(0,i), h, True); shade_cell(corr.cell(0,i))
rows = [
    ("01", "Catatan transaksi", "Ekstraksi nilai", "0.45 BTC"),
    ("02", "TXID / cache wallet", "Korelasi transaksi", "17SkEw2md5avVNyYgj6RiXuQKNwkXaxFyQ"),
    ("03", "beacon_cfg.dat + PCAP", "Base64 → XOR → DNS correlation", "NW-7739|185.220.101.47"),
    ("04", "Citra + EXIF GPS", "Metadata analysis", "DANAU BERATAN"),
]
for row in rows:
    cells = corr.add_row().cells
    for i,v in enumerate(row):
        set_cell_text(cells[i], v)

doc.add_heading("9. Chain of Custody / Penanganan Evidence", level=1)
doc.add_paragraph(
    "Dokumentasi berikut merupakan format chain-of-custody untuk keperluan laporan challenge. "
    "Hash kriptografis sumber sebaiknya ditambahkan apabila image/disk evidence asli tersedia."
)
cot = doc.add_table(rows=1, cols=5)
cot.style = "Table Grid"
cot.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(["No.", "Evidence", "Sumber", "Tindakan", "Status"]):
    set_cell_text(cot.cell(0,i), h, True); shade_cell(cot.cell(0,i))
for i,row in enumerate([
    ("1","Artefak Node 01","File challenge","Analisis read-only","Teranalisis"),
    ("2","Catatan transaksi Node 02","Cache wallet","Ekstraksi TXID/address","Teranalisis"),
    ("3","beacon_cfg.dat","File challenge","Base64 + XOR analysis","Teranalisis"),
    ("4","PCAP Node 03","Capture jaringan","DNS correlation","Teranalisis"),
    ("5","Image Node 04","File gambar","EXIF/GPS extraction","Teranalisis"),
],1):
    cells=cot.add_row().cells
    for j,v in enumerate((i,)+row):
        set_cell_text(cells[j], v)

doc.add_heading("10. Kesimpulan", level=1)
doc.add_paragraph(
    "Berdasarkan analisis dan korelasi evidence, seluruh node dapat diselesaikan dengan "
    "metode forensik yang berbeda. Node 03 merupakan kasus yang membutuhkan korelasi "
    "lintas-artifact: beacon configuration menghasilkan kode operasi dan hostname, "
    "sedangkan PCAP menyediakan resolusi hostname ke IP callback. Node 04 divalidasi "
    "melalui metadata GPS pada citra."
)
p = doc.add_paragraph()
p.add_run("Rekapitulasi jawaban final:").bold = True
for ans in [
    "Node 01 → 0.45 BTC",
    "Node 02 → 17SkEw2md5avVNyYgj6RiXuQKNwkXaxFyQ",
    "Node 03 → NW-7739|185.220.101.47",
    "Node 04 → DANAU BERATAN",
]:
    doc.add_paragraph(ans, style="List Bullet")

doc.add_heading("11. Catatan Forensik", level=1)
doc.add_paragraph(
    "Laporan ini ditujukan sebagai dokumentasi analisis challenge/CTF. Beberapa detail "
    "chain-of-custody seperti hash SHA-256 image disk, identitas examiner, waktu akuisisi, "
    "dan nomor evidence harus diisi dari proses akuisisi asli apabila laporan digunakan "
    "sebagai dokumen investigasi formal."
)

# Footer
for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Digital Forensic Report — Chapter 4").font.size = Pt(8)

doc.save(out)
print(out)
